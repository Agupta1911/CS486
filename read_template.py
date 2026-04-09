import docx

def read_docx(filename):
    doc = docx.Document(filename)
    for i, paragraph in enumerate(doc.paragraphs):
        print(f"P[{i}]: {paragraph.text}")
    print("--- TABLES ---")
    for t, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                print(f"T[{t}] R[{r}] C[{c}]: {cell.text}")

if __name__ == '__main__':
    read_docx('HCAI_Milestone1_Template.docx')
