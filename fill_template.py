import docx

doc = docx.Document('HCAI_Milestone1_Template.docx')

doc.paragraphs[9].text = "The primary audience for this experimental chatbot prototype consists of retail investors, amateur day traders, and finance students. These individuals are actively looking to analyze stock market data without needing advanced financial terminals or complex research software."
doc.paragraphs[11].text = "This demographic often struggles with information overload when researching stocks, as they must typically synthesize quantitative data (like price and volume) with qualitative data (like recent news) from multiple different sources. By providing a conversational interface that centralizes this information, we empower them to make more informed and timely financial decisions."
doc.paragraphs[13].text = "Knowledge synthesis is crucial for this audience because stock market movements are driven by both hard numbers and public sentiment. Being able to quickly request and receive a cohesive summary of a stock's history alongside contextual news allows these users to perform rapid, effective analysis that is otherwise time-consuming and fragmented."

doc.tables[0].cell(0, 1).text = "Aarav, Axel"
doc.tables[1].cell(0, 1).text = "Knowledge synthesis and summarization"
doc.tables[1].cell(1, 1).text = "N/A"

doc.tables[2].cell(1, 1).text = "A retail investor wants to quickly understand the historical price trends, trading volume, and recent news for a specific stock by providing the ticker symbol to the chatbot, which then retrieves and synthesizes data directly from Yahoo Finance."
doc.tables[2].cell(2, 1).text = "A student in a finance class needs to summarize recent market events impacting the tech sector for a class presentation, and asks the chatbot to aggregate news and historical metrics to provide a comprehensive, easy-to-read overview."
doc.tables[2].cell(3, 1).text = "An amateur day trader wants to check if a sudden drop in a stock's volume corresponds to any recent controversies or earnings reports, using the chatbot to fetch both quantitative financial data and qualitative news summaries simultaneously."

doc.save('HCAI_Milestone1_Completed.docx')
print("Successfully saved to HCAI_Milestone1_Completed.docx")
